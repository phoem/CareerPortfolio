#!/usr/bin/env python3
"""Generate DOCX and PDF artifacts from canonical Markdown sources."""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]


@dataclass(frozen=True)
class SourceSpec:
    source: Path
    artifact_type: str
    design_name: str
    one_shot_config: Path | None = None


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def default_design_name() -> str:
    return str(load_json(ROOT / "designs" / "default.json")["design"])


def resolve_application_design(directory: Path, artifact_type: str) -> tuple[str, Path | None]:
    for filename, one_shot in (("DESIGN_NEXT.json", True), ("DESIGN.json", False)):
        path = directory / filename
        if not path.exists():
            continue
        config = load_json(path)
        applies_to = config.get("applies_to", ["resume", "cover_letter"])
        if artifact_type in applies_to:
            return str(config["design"]), path if one_shot else None
    return default_design_name(), None


def design_config(name: str, artifact_type: str) -> dict:
    path = ROOT / "designs" / name / "design.json"
    if not path.exists():
        raise FileNotFoundError(f"Unknown design '{name}': {path} does not exist")
    config = load_json(path)
    section = "cover_letter" if artifact_type == "cover_letter" else "resume"
    if section not in config:
        raise ValueError(f"Design '{name}' does not define a {section} section")
    return {"font": config.get("font", "Arial"), **config[section]}


def discover_sources() -> list[SourceSpec]:
    sources: list[SourceSpec] = []
    default = default_design_name()
    for source in sorted((ROOT / "generic").glob("*.md")):
        sources.append(SourceSpec(source, "resume", default))

    for manifest_path in sorted((ROOT / "applications").glob("**/APPLICATION.json")):
        directory = manifest_path.parent
        manifest = load_json(manifest_path)
        for artifact in manifest.get("artifacts", []):
            source = directory / artifact["source"]
            artifact_type = str(artifact["type"])
            if not source.exists():
                raise FileNotFoundError(f"Manifest source does not exist: {source}")
            design_name, one_shot = resolve_application_design(directory, artifact_type)
            sources.append(SourceSpec(source, artifact_type, design_name, one_shot))
    return sources


def relative(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()


def select_sources(
    sources: list[SourceSpec],
    *,
    full_rebuild: bool,
    requested_sources: set[str],
    changed_paths: set[str],
) -> tuple[list[SourceSpec], list[str]]:
    """Select only artifacts explicitly affected by this run.

    Shared generator and design changes deliberately do not select every resume.
    They return a recommendation so the user can decide whether to dispatch a
    manual full rebuild.
    """
    by_path = {relative(spec.source): spec for spec in sources}
    selected: set[str] = set()
    recommendations: list[str] = []

    if full_rebuild:
        return sources, recommendations

    for requested in requested_sources:
        normalized = Path(requested.strip()).as_posix().lstrip("./")
        if normalized not in by_path:
            raise ValueError(f"Requested source is not a generated artifact source: {requested}")
        selected.add(normalized)

    for changed in changed_paths:
        normalized = Path(changed.strip()).as_posix().lstrip("./")
        if not normalized:
            continue
        if normalized in by_path:
            selected.add(normalized)
            continue

        name = Path(normalized).name
        if normalized.startswith("applications/") and name in {
            "APPLICATION.json",
            "DESIGN.json",
            "DESIGN_NEXT.json",
        }:
            directory = Path(normalized).parent.as_posix()
            selected.update(
                source_path
                for source_path in by_path
                if Path(source_path).parent.as_posix() == directory
            )
            continue

        if normalized == "scripts/generate_resume_artifacts.py" or normalized.startswith(
            "designs/"
        ):
            recommendations.append(
                f"{normalized} changed; consider a manual full rebuild after reviewing the impact."
            )

    return [by_path[path] for path in sorted(selected)], recommendations


def add_bottom_border(
    paragraph,
    *,
    color: str = "666666",
    size: int = 5,
    space: int = 1,
) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    properties.append(borders)


def add_template_horizontal_rule(document, config: dict) -> None:
    """Add the compact gray bar used by the distilled DEH reference template."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(
        float(config.get("left_indent_inches", -0.0625))
    )
    paragraph.paragraph_format.line_spacing = float(config.get("line_spacing", 0.4))
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    height = float(config.get("height_points", 1.5))
    color = str(config.get("color", "A0A0A0")).lstrip("#")
    run = paragraph.add_run()
    run._r.append(
        parse_xml(
            f'<w:pict {nsdecls("w")} '
            f'xmlns:v="urn:schemas-microsoft-com:vml" '
            f'xmlns:o="urn:schemas-microsoft-com:office:office">'
            f'<v:rect style="width:0.0pt;height:{height}pt" '
            f'o:hr="t" o:hrstd="t" o:hralign="center" '
            f'fillcolor="#{color}" stroked="f"/>'
            f'</w:pict>'
        )
    )


def add_inline_markup(paragraph, text: str) -> None:
    for part in re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            paragraph.add_run(part[1:-1])
        else:
            paragraph.add_run(part)


def markdown_to_docx(spec: SourceSpec, destination: Path) -> None:
    style = design_config(spec.design_name, spec.artifact_type)
    is_letter = spec.artifact_type == "cover_letter"
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(float(style["top_margin_inches"]))
    section.bottom_margin = Inches(float(style["bottom_margin_inches"]))
    section.left_margin = Inches(float(style["left_margin_inches"]))
    section.right_margin = Inches(float(style["right_margin_inches"]))

    normal = document.styles["Normal"]
    normal.font.name = style["font"]
    normal.font.size = Pt(float(style["body_font_points"]))
    normal.paragraph_format.space_after = Pt(
        float(style.get("body_space_after_points", 5 if is_letter else 1.2))
    )
    normal.paragraph_format.line_spacing = float(style["line_spacing"])

    heading_defaults = (("Heading 1", 11.2), ("Heading 2", 10.0), ("Heading 3", 9.2))
    for style_name, default_size in heading_defaults:
        level = style_name[-1]
        heading_style = document.styles[style_name]
        heading_style.font.name = style["font"]
        heading_style.font.bold = True
        if "heading_color" in style:
            heading_style.font.color.rgb = RGBColor.from_string(
                str(style["heading_color"]).lstrip("#")
            )
        heading_style.font.size = Pt(
            float(style.get(f"heading_{level}_font_points", default_size))
        )
        heading_style.paragraph_format.space_before = Pt(
            float(style.get(f"heading_{level}_space_before_points", 3))
        )
        heading_style.paragraph_format.space_after = Pt(
            float(style.get(f"heading_{level}_space_after_points", 1))
        )
        heading_style.paragraph_format.keep_with_next = True

    for line in spec.source.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(
                float(style.get("title_space_after_points", 0))
            )
            title = line[2:]
            if title.startswith("Jordan Newman -"):
                title = "JORDAN NEWMAN"
            run = paragraph.add_run(title)
            run.bold = True
            run.font.size = Pt(float(style["title_font_points"]))
        elif line.startswith("## "):
            paragraph = document.add_paragraph(style="Heading 1")
            paragraph.add_run(line[3:].upper())
            rule = style.get("section_rule", {})
            if rule.get("style") == "template-bar":
                add_template_horizontal_rule(document, rule)
            else:
                add_bottom_border(
                    paragraph,
                    color=str(rule.get("color", "666666")).lstrip("#"),
                    size=int(rule.get("size", 5)),
                    space=int(rule.get("space", 1)),
                )
        elif line.startswith("### "):
            paragraph = document.add_paragraph(style="Heading 2")
            paragraph.add_run(line[4:])
        elif line.startswith("- "):
            use_real_bullets = bool(style.get("use_real_bullets", False))
            paragraph = document.add_paragraph(
                style="List Bullet" if use_real_bullets else None
            )
            paragraph.paragraph_format.left_indent = Inches(
                float(style.get("bullet_left_indent_inches", 0.16))
            )
            paragraph.paragraph_format.first_line_indent = Inches(
                float(style.get("bullet_first_line_indent_inches", -0.12))
            )
            paragraph.paragraph_format.space_after = Pt(
                float(style.get("bullet_space_after_points", 0.8))
            )
            if not use_real_bullets:
                paragraph.add_run("• ").bold = True
            add_inline_markup(paragraph, line[2:])
        elif line.startswith("**") and line.endswith("**"):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if is_letter else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(line[2:-2])
            run.bold = True
            run.font.size = Pt(float(style.get("emphasis_font_points", 10.2)))
        elif line.startswith("*") and line.endswith("*"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            # Keep resume job metadata (location and dates) with the first
            # accomplishment bullet so a role heading is not orphaned at a
            # page boundary. Cover-letter emphasis remains free-flowing.
            paragraph.paragraph_format.keep_with_next = not is_letter
            paragraph.add_run(line[1:-1]).italic = True
        else:
            paragraph = document.add_paragraph()
            if not is_letter and line.startswith("Marlboro, NJ"):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(2)
            add_inline_markup(paragraph, line)

    if bool(style.get("show_footer", True)):
        footer = document.sections[0].footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Jordan Newman").font.size = Pt(
            float(style.get("footer_font_points", 7))
        )
    document.save(destination)


def convert_to_pdf(docx_path: Path) -> None:
    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    if libreoffice:
        subprocess.run(
            [
                libreoffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(docx_path.parent),
                str(docx_path),
            ],
            check=True,
        )
        return

    if os.name == "nt":
        powershell = shutil.which("powershell.exe") or shutil.which("powershell")
        word_converter = ROOT / "scripts" / "convert_docx_to_pdf_with_word.ps1"
        if powershell and word_converter.exists():
            subprocess.run(
                [
                    powershell,
                    "-NoProfile",
                    "-NonInteractive",
                    "-ExecutionPolicy",
                    "Bypass",
                    "-File",
                    str(word_converter),
                    "-DocxPath",
                    str(docx_path.resolve()),
                ],
                check=True,
            )
            return

    raise RuntimeError(
        "PDF conversion requires LibreOffice/soffice, or Microsoft Word on Windows."
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--full-rebuild",
        action="store_true",
        help="Rebuild every generated artifact. Intended only for explicit manual use.",
    )
    parser.add_argument(
        "--source",
        action="append",
        default=[],
        help="Repository-relative Markdown artifact source to rebuild (repeatable).",
    )
    parser.add_argument(
        "--changed-paths-file",
        type=Path,
        help="File containing repository-relative changed paths, one per line.",
    )
    parser.add_argument(
        "--rebuilt-manifest",
        type=Path,
        required=True,
        help="Write the exact rebuilt artifact set and rebuild recommendations as JSON.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    changed_paths: set[str] = set()
    if args.changed_paths_file and args.changed_paths_file.exists():
        changed_paths = {
            line.strip()
            for line in args.changed_paths_file.read_text(encoding="utf-8").splitlines()
            if line.strip()
        }

    sources, recommendations = select_sources(
        discover_sources(),
        full_rebuild=args.full_rebuild,
        requested_sources=set(args.source),
        changed_paths=changed_paths,
    )
    consumed_one_shot: set[Path] = set()
    rebuilt: list[dict[str, str]] = []
    for spec in sources:
        destination = spec.source.with_suffix(".docx")
        markdown_to_docx(spec, destination)
        convert_to_pdf(destination)
        if spec.one_shot_config:
            consumed_one_shot.add(spec.one_shot_config)
        print(
            f"Generated {destination.relative_to(ROOT)} and "
            f"{destination.with_suffix('.pdf').relative_to(ROOT)} "
            f"using design {spec.design_name}"
        )
        rebuilt.append(
            {
                "source": relative(spec.source),
                "artifact_type": spec.artifact_type,
                "docx": relative(destination),
                "pdf": relative(destination.with_suffix(".pdf")),
                "design": spec.design_name,
            }
        )
    for path in consumed_one_shot:
        path.unlink(missing_ok=True)
        print(f"Consumed one-build design override {path.relative_to(ROOT)}")

    result = {
        "full_rebuild": bool(args.full_rebuild),
        "rebuilt": rebuilt,
        "full_rebuild_recommended": bool(recommendations),
        "recommendations": sorted(set(recommendations)),
    }
    args.rebuilt_manifest.parent.mkdir(parents=True, exist_ok=True)
    args.rebuilt_manifest.write_text(json.dumps(result, indent=2) + "\n", encoding="utf-8")
    if not rebuilt:
        print("No resume artifacts were selected for rebuilding.")
    for recommendation in result["recommendations"]:
        print(f"Recommendation: {recommendation}")


if __name__ == "__main__":
    main()
